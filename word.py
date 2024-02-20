from docx import Document
import pathlib
from tkinter import messagebox


class GenerateWord:
    try:
        def replace_text_in_docx(self, replacements, path):
            doc = Document(path)
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value)
            doc.save(f"{pathlib.Path().resolve()}\\gеnerated-files\\{replacements['<<name>>']}.docx")
            messagebox.showinfo("Инфо", f"файлът с допълнителното споразумение на {replacements['<<name>>']} е генераран!")
    except SyntaxError as SErr:
        messagebox.showerror("Error", SErr)
        raise SErr

